# 项目：我的工具箱
# 模块：生成 Word 模块
# 作者：黄涛
# License:GPL
# Email:huangtao.sh@icloud.com
# 创建：2023-04-21 23:35
# 修订：2025-09-20 10:22 类型检查优化

from docx import Document as _Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Pt
from docx.text.font import Font
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.styles.style import ParagraphStyle

from orange import Path, datetime, today, wlen
from typing import Union, Optional


from .utils.hz import Ordinal


def Remove(elem: Union[Paragraph, Table]):
    "删除指定的元素，包括：段落、表格、图片等"
    return elem._element.getparent().remove(elem._element)


def SetFontName(font: Font, value: str):
    """增加中文字体的设置"""
    rPr = font._element.get_or_add_rPr()
    rPr.rFonts_ascii = value
    rPr.rFonts_hAnsi = value
    rFonts = rPr.rFonts
    assert rFonts is not None
    rFonts.set(qn("w:eastAsia"), value)


def insert_paragraph_after(
    self: Paragraph, text: str = "", style: Optional[ParagraphStyle] = None
) -> Paragraph:
    "在指定的段落后面插入段落"
    "看不太懂，从网上抄的，貌似可以正常工作"
    new_p = OxmlElement("w:p")
    self._p.addnext(new_p)
    new_para = Paragraph(new_p, self._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


setattr(
    Font, "name", Font.name.setter(SetFontName)
)  # 替换掉原来的设置字体，以便支持中文字体
setattr(Paragraph, "insert_paragraph_after", insert_paragraph_after)
# Paragraph.insert_paragraph_after = (
#    insert_paragraph_after  # 为段落对象增加在段后增加段落的功能
# )


def set_table_border(table, auto=True, **kwargs):
    """
    设置表格的边框
    用法:
    set_table_border(
        table,
    top={"sz": 12, "val": "single", "color": "#FF0000"},
    bottom={"sz": 12, "color": "#00FF00", "val": "single"},
    left={"sz": 24, "val": "dashed"},
    right={"sz": 12, "val": "dashed"},
    )
    """
    # default = {"sz": 1, "val": "single", "color": "#000000", "space": "0"}
    borders = OxmlElement("w:tblBorders")
    for tag in ("bottom", "top", "left", "right", "insideV", "insideH"):
        edge_data = kwargs.get(tag)
        if edge_data:
            any_border = OxmlElement(f"w:{tag}")
            for key in ["sz", "val", "color", "space", "shadow"]:
                assert isinstance(edge_data, dict)
                if key in edge_data:
                    assert isinstance(any_border, BaseOxmlElement)
                    any_border.set(qn(f"w:{key}"), str(edge_data[key]))
                    borders.append(any_border)
    table._tbl.tblPr.append(borders)


class Document:
    """
    标准公文格式
    """

    def __init__(self):
        self.document = _Document()
        self.heading = [Ordinal(), Ordinal(), Ordinal(), Ordinal()]
        # section = self.document.sections[0]
        # section.top_margin = Cm(3.7)
        # section.bottom_margin = Cm(3.5)
        # section.left_margin = Cm(2.8)
        # section.right_margin = Cm(2.6)

        zw = self.document.styles["Normal"]
        assert isinstance(zw, ParagraphStyle)
        zw.font.size = Pt(16)
        zw.font.name = "仿宋_GB2312"
        zw.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        zw.paragraph_format.first_line_indent = Pt(32)
        zw.paragraph_format.space_after = Pt(0)
        zw.paragraph_format.line_spacing = 1.5
        # zw.keep_together = False

        gwbt = self.document.styles.add_style("公文标题", WD_STYLE_TYPE.PARAGRAPH)
        assert isinstance(gwbt, ParagraphStyle)
        gwbt.font.size = Pt(22)
        gwbt.font.name = "方正小标宋简体"
        gwbt.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        gwbt.paragraph_format.space_after = Pt(0)
        gwbt.paragraph_format.line_spacing = 1

        zw = self.document.styles.add_style("一级标题", WD_STYLE_TYPE.PARAGRAPH)
        assert isinstance(zw, ParagraphStyle)
        zw.font.size = Pt(16)
        zw.font.name = "黑体"
        zw.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        zw.paragraph_format.first_line_indent = Pt(32)
        zw.paragraph_format.space_after = Pt(0)
        zw.paragraph_format.line_spacing = 1.5

        zw = self.document.styles.add_style("二级标题", WD_STYLE_TYPE.PARAGRAPH)
        assert isinstance(zw, ParagraphStyle)
        zw.font.size = Pt(16)
        zw.font.name = "楷体"
        zw.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        zw.paragraph_format.first_line_indent = Pt(32)
        zw.paragraph_format.space_after = Pt(0)
        zw.paragraph_format.line_spacing = 1.5

        zw = self.document.styles.add_style("文号", WD_STYLE_TYPE.PARAGRAPH)
        assert isinstance(zw, ParagraphStyle)
        zw.font.size = Pt(14)
        zw.font.name = "仿宋_GB2312"
        zw.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        zw.paragraph_format.space_after = Pt(0)
        zw.paragraph_format.line_spacing = 1

    def save(self, path: Union[str, Path]):
        "保存到磁盘文件"
        self.document.save(str(Path(path)))

    def add_title(self, title: str):
        "添加标题"
        return self.document.add_paragraph(title, "公文标题")

    def add_wenhao(self, wenhao: str):
        "添加文号"
        return self.document.add_paragraph(wenhao, "文号")

    def add_zsjg(self, jgmc: str):
        "添加主送机构"
        if not jgmc.endswith("："):
            jgmc += "："
        p = self.document.add_paragraph(jgmc, "Normal")
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.first_line_indent = Pt(0)

    def add_para(self, lines):
        "添加正文"
        for p in lines.splitlines():
            self.document.add_paragraph(p, "Normal")

    def add_heading(self, head, level=1):
        "添加标题"
        if not 0 < level < 5:
            raise Exception(f"小标题 {level} 错误")
        if level == 1:
            style = "一级标题"
        elif level == 2:
            style = "二级标题"
        else:
            style = "Normal"
        self.document.add_paragraph(head, style)

    def add_fwjg(self, fwjg: str, rq: str = ""):
        "添加发文机构"
        if not rq:
            rq = today()
        rq = datetime(rq).format("%x")
        self.add_para("")
        p = self.document.add_paragraph(fwjg, "Normal")
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.paragraph_format.right_indent = Pt(
            ((wlen(rq) + 3) / 2 + 8 - wlen(fwjg) / 2) * 8
        )
        p = self.document.add_paragraph(rq, "Normal")
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.paragraph_format.right_indent = Pt(16 * 4)
